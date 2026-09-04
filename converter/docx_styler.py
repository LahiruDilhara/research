import os
import sys
import docx
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL

def style_thesis_docx(docx_path):
    print(f'Loading {docx_path}...')
    doc = docx.Document(docx_path)
    
    # 1. Page Margins (1 inch = 1440 dxa)
    for s in doc.sections:
        s.top_margin = Inches(1.0)
        s.bottom_margin = Inches(1.0)
        s.left_margin = Inches(1.0)
        s.right_margin = Inches(1.0)
        s.page_width = Inches(8.5)
        s.page_height = Inches(11.0)
        
    TOTAL_WIDTH = 9360  # 6.5 in printable area in dxa
    
    # 2. Typography & Headings
    for p in doc.paragraphs:
        st = p.style.name if p.style else ''
        t = p.text.strip()
        
        if 'Title' in st:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(24)
            p.paragraph_format.space_after = Pt(12)
            for r in p.runs:
                r.font.name = 'Calibri'
                r.font.size = Pt(22)
                r.font.bold = True
                r.font.color.rgb = RGBColor(0x0F, 0x29, 0x42)
                
        elif 'Heading 1' in st:
            p.paragraph_format.space_before = Pt(22)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.keep_with_next = True
            if t and not t.startswith('List of Abbreviations') and any(str(i) in t for i in range(1, 8)):
                p.paragraph_format.page_break_before = True
            elif 'References' in t:
                p.paragraph_format.page_break_before = True
                
            for r in p.runs:
                r.font.name = 'Calibri'
                r.font.size = Pt(17)
                r.font.bold = True
                r.font.color.rgb = RGBColor(0x0F, 0x29, 0x42)
                
        elif 'Heading 2' in st:
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            for r in p.runs:
                r.font.name = 'Calibri'
                r.font.size = Pt(13.5)
                r.font.bold = True
                r.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
                
        elif 'Heading 3' in st:
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.keep_with_next = True
            for r in p.runs:
                r.font.name = 'Calibri'
                r.font.size = Pt(11.5)
                r.font.bold = True
                r.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
                
        elif 'Caption' in st or t.startswith('Figure ') or t.startswith('Table '):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(10)
            for r in p.runs:
                r.font.name = 'Calibri'
                r.font.size = Pt(10)
                r.font.italic = True
                r.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
                
        else:
            p.paragraph_format.space_after = Pt(5)
            p.paragraph_format.line_spacing = 1.15
            for r in p.runs:
                if not r.font.name:
                    r.font.name = 'Calibri'
                if not r.font.size:
                    r.font.size = Pt(11)
                if not r.font.color.rgb:
                    r.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
                    
    # 3. Format Tables
    print(f'Formatting {len(doc.tables)} tables...')
    for tbl in doc.tables:
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tblPr = tbl._tbl.tblPr
        
        # Set Table Width
        tblW = tblPr.find(qn('w:tblW'))
        if tblW is None:
            tblW = OxmlElement('w:tblW')
            tblPr.append(tblW)
        tblW.set(qn('w:w'), str(TOTAL_WIDTH))
        tblW.set(qn('w:type'), 'dxa')
        
        # Borders
        borders_xml = (
            f'<w:tblBorders {nsdecls("w")}>'
            '<w:top w:val="single" w:sz="12" w:space="0" w:color="1B365D"/>'
            '<w:left w:val="none"/>'
            '<w:bottom w:val="single" w:sz="12" w:space="0" w:color="1B365D"/>'
            '<w:right w:val="none"/>'
            '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>'
            '<w:insideV w:val="none"/>'
            '</w:tblBorders>'
        )
        old_b = tblPr.find(qn('w:tblBorders'))
        if old_b is not None:
            tblPr.remove(old_b)
        tblPr.append(parse_xml(borders_xml))
        
        ncols = len(tbl.columns)
        if ncols == 0:
            continue
            
        if ncols == 2:
            col_widths = [2340, 7020]
        elif ncols == 3:
            col_widths = [1680, 2520, 5160]
        elif ncols == 4:
            col_widths = [1872, 2340, 2808, 2340]
        elif ncols == 5:
            col_widths = [1872] * 5
        elif ncols >= 7:
            base_w = TOTAL_WIDTH // ncols
            col_widths = [base_w] * ncols
            col_widths[-1] += TOTAL_WIDTH - sum(col_widths)
        else:
            base_w = TOTAL_WIDTH // ncols
            col_widths = [base_w] * ncols
            col_widths[-1] += TOTAL_WIDTH - sum(col_widths)
            
        is_dense = (ncols >= 7)
        
        for r_idx, row in enumerate(tbl.rows):
            trPr = row._tr.get_or_add_trPr()
            cantSplit = OxmlElement('w:cantSplit')
            trPr.append(cantSplit)
            
            is_header = (r_idx == 0)
            if is_header:
                tblHeader = OxmlElement('w:tblHeader')
                trPr.append(tblHeader)
                
            for c_idx, cell in enumerate(row.cells):
                tcPr = cell._tc.get_or_add_tcPr()
                
                if c_idx < len(col_widths):
                    tcW = tcPr.find(qn('w:tcW'))
                    if tcW is None:
                        tcW = OxmlElement('w:tcW')
                        tcPr.append(tcW)
                    tcW.set(qn('w:w'), str(col_widths[c_idx]))
                    tcW.set(qn('w:type'), 'dxa')
                    
                tcMar = OxmlElement('w:tcMar')
                for m_name, val in [('top', 120), ('bottom', 120), ('left', 160), ('right', 160)]:
                    node = OxmlElement(f'w:{m_name}')
                    node.set(qn('w:w'), str(val))
                    node.set(qn('w:type'), 'dxa')
                    tcMar.append(node)
                tcPr.append(tcMar)
                
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                if is_header:
                    shd_xml = f'<w:shd {nsdecls("w")} w:fill="1B365D"/>'
                    tcPr.append(parse_xml(shd_xml))
                    for p in cell.paragraphs:
                        p.paragraph_format.space_before = Pt(2)
                        p.paragraph_format.space_after = Pt(2)
                        p.paragraph_format.line_spacing = 1.05
                        for r in p.runs:
                            r.font.name = 'Calibri'
                            r.font.size = Pt(8.5 if is_dense else 9.5)
                            r.font.bold = True
                            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                else:
                    bg_col = 'F8FAFC' if (r_idx % 2 == 0) else 'FFFFFF'
                    if bg_col != 'FFFFFF':
                        shd_xml = f'<w:shd {nsdecls("w")} w:fill="{bg_col}"/>'
                        tcPr.append(parse_xml(shd_xml))
                        
                    for p in cell.paragraphs:
                        p.paragraph_format.space_before = Pt(1.5)
                        p.paragraph_format.space_after = Pt(1.5)
                        p.paragraph_format.line_spacing = 1.08
                        for r in p.runs:
                            r.font.name = 'Calibri'
                            r.font.size = Pt(8.0 if is_dense else 9.5)
                            r.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
                            
    doc.save(docx_path)
    print(f'Successfully styled {docx_path}!')

if __name__ == '__main__':
    target = sys.argv[1] if len(sys.argv) > 1 else 'thesis.docx'
    style_thesis_docx(target)
